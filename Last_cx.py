


import requests
import re                                                                   # 导入re模块

from playsound import playsound
from bs4 import BeautifulSoup
from docx import Document                                      # 导入Document模块
from docx.shared import Pt                                       # 导入Pt模块



def crawl_urls1(number):                                       # 获取首页前二十篇文章。
        base_url = 'https://news.cri.cn/guojiruiping'
        response = requests.get(base_url)
        html = response.text

        soup = BeautifulSoup(html, 'html.parser')         # 解析网页内容

        # 定位并提取链接
        links = soup.select('div.sftx-list.more-list a')
        urls = [link['href'] for link in links if link.has_attr('href')]
        return urls


def crawl_urls2(i):                                               # 抓取 20 以后的文章
        i = str(i)
        base_url ='https://news.cri.cn/inc/08b152bb-56a6-4ba4-9b60-8761fa1a1568-'

        next_url = base_url + i +'.inc'                         #拼接url

        response = requests.get(next_url)
        html = response.text

        soup = BeautifulSoup(html, 'html.parser')                     # 解析网页内容

         # 定位并提取链接
        links = soup.select('div.sftx-list.more-list a')
        urls = list(set([link['href'] for link in links if link.has_attr('href')]))
        return urls



def crawl_articles(url):                                             # 获取指定URL的文本

    if url is None or url == "":  # 判断url是否为None或空字符串
        print("此url为空，已跳过")
        return  # 跳出函数
    else:

        url = url.replace("//news.cri.cn", "")
        # 获取网页内容
        response = requests.get('https://news.cri.cn' + url)
        html = response.text

        # 解析网页内容
        soup = BeautifulSoup(html, 'html.parser')

        try:
            title = soup.find('div', class_="list-title")            # 尝试定位并提取标题
            title_text = title.get_text()                                 # 尝试获取标题文本
        except AttributeError:                                          # 如果失败，捕获异常
            title_text = '网址错误'                                     # 返回一个空值
            print('网页没有标题或者标题无法获取')
            print('https://news.cri.cn' + url)                        # 提示用户网页没有标题或者标题无法获取

        # 定位并提取时间
        time = soup.find('div', class_="list-brief")
        time_article_ = time.get_text()
        time_article = time_article_[:10]


        # 定位并提取内容文本
        article = soup.find('div', class_="list-abody abody")
        article_text = article.get_text()

        document = Document()                                         # 创建一个新的 Word 文档

        # 设置样式
        style = document.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)

        # 将文章的文本内容添加到文档中
        document.add_heading(title_text)
        document.add_paragraph(article_text)

        # 设置文件名
        filename = re.sub(r'[\/\\\:\*\?\"\<\>\|]', '',time_article +  title_text)

        # 保存Word文档
        document.save(filename + '.docx')


# def music_():
#     print('需要一首音乐？\n')
#     answer = input('请回答需要或不需要')
#     while True:
#         # try:
#         if answer is '需要':
#             print('1.光芒，2.归途有风，3.玫瑰少年')
#             music =  input('请选择一首歌,输入对应序号')
#             if music is '1' :
#                 playsound('光芒.mp3')
#             elif music is '2' :
#                 playsound('归途有风.mp3')
#             elif music is '3' :
#                 playsound('玫瑰少年.mp3')
#             else:
#                 print('请输入正确序号')
#
#
#
#         elif answer is '不需要':
#             break
#
#         else:
#             return
#
#         # else :
#         #     print('请正确输入')








if __name__ == '__main__':
    while True:                                                                 # 无限循环
        try:

            print('你好！我是一个国际锐评文章爬取程序\n')
            number = int(input('请问你需要获取多少篇文章？\n（请输入20的整数倍，例如20,40,60,80......）\n'))  # 获取用户输入的整数
            # starry_sky()
            print('请稍等......')

            if number <= 20:
                urls = crawl_urls1(number)
                for url in urls:
                    crawl_articles(url)
                print('已完成！')

            else:
                if number%20 != 0 :
                    print('你干嘛~~~~哎~~~哟 \n（请输入20的整数倍）')
                    playsound('niganma.mp3')  # 播放本地文件
                    # playsound('https://example.com/sample.mp3')  # 播放网络文件
                # 获取前二十个
                urls = crawl_urls1(number)
                for url in urls:
                    crawl_articles(url)

                page = number // 20
                for i in range(2, page + 1):
                    urls = crawl_urls2(i)

                    for url in urls:
                        crawl_articles(url)
                    urls.clear()

                print('已完成！')
            break  # 如果成功，跳出循环

        except ValueError:  # 如果失败，捕获异常
            print('请重新输入！')  # 提示用户重新输入



